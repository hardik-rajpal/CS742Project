from paraphraser import Rephraser
import zlib
import re
import inspect
import uuid
import random
from docx2pdf import convert
from docx import Document
import math
import spacy
import argparse
SIMILARITY_THRESHOLD=0.8
SERVER_KEY = 'howard_roark'
DEFAULT_DOC_VARS = [
  ["Test Document"],
  [
    "Online Education; it\u2019s a curse for some students and a blessing for others.",
    "Online education can be challenging for some students, while it offers benefits for others.",
    "The online learning experience can be overwhelming for some students, but it is advantageous for others.",
    "Some students find online education to be a burden, while others find it beneficial."
  ],
  [
    "The part that pleases me the most is the ease with which we can switch classrooms.",
    "The aspect of my job that I enjoy the most is the convenience of being able to switch between classrooms easily.",
    "The thing that brings me the most satisfaction in my work is the ability to move between classrooms without any hassle.",
    "The part of my job that I find the most enjoyable is the flexibility to switch between classrooms with ease."
  ],
  [
    "The only way I am gonna be happy if the campus reopens is if the lecture halls are a Chrome tab away.",
    "The only way I can be content if the campus reopens is if the lecture halls are accessible through a Chrome tab.",
    "If the lecture halls are available through a Chrome tab, then I will be satisfied with the campus reopening.",
    "The only way I can feel happy if the campus reopens is if the lecture halls are accessible through a Chrome tab."
  ],
  [
    "The idea of a \"flipped-classroom\" model, where the students watch the pre-recorded lectures before having a live class, really relaxes the time constraints faced in traditional schooling.",
    "The concept of a flipped classroom, where students watch pre-recorded lectures before having live classes, is a great way to reduce time constraints in traditional schooling.",
    "The idea of a flipped classroom, where students watch pre-recorded lectures before having live classes, can help alleviate the pressure of traditional schooling.",
    "The concept of a flipped classroom, where students watch pre-recorded lectures before having live classes, is an effective way to reduce time constraints in traditional schooling."
  ],
  [
    "As advocated by Sal Khan from Khan Academy, the whole process can become perfection-centred rather than the usual complete-the-syllabus-on-time-centred.",
    "As suggested by Sal Khan from Khan Academy, the entire process can be transformed into a perfection-centered approach instead of the traditional \"complete the syllabus on time\" mindset.",
    "Sal Khan's approach to education emphasizes a perfection-centered approach that focuses on developing skills and fostering creativity rather than simply completing tasks within a set timeframe.",
    "Sal Khan's philosophy of education places emphasis on a perfection-centered approach, which prioritizes cultivating skills and fostering creativity over adhering to strict deadlines or completing tasks within a specific time frame."
  ],
  [
    "This fluidity of the schedule permits one to partition their time to maximise their efficiency and clarity in a concept; a point sadly overlooked in the conventional and offline teaching methods.",
    "The flexibility of the schedule allows individuals to allocate their time effectively, maximizing their efficiency and clarity in understanding concepts, which is often neglected in traditional and offline teaching methods.",
    "The fluidity of the schedule enables individuals to manage their time efficiently and gain a deeper understanding of concepts, which is often overlooked in conventional and offline teaching methods.",
    "The flexibility of the schedule allows individuals to allocate their time effectively, maximizing their efficiency and clarity in understanding concepts, which is often overlooked in traditional and offline teaching methods."
  ],
  [
    "These ideas being the bright side of the issue; there are screen-sharing mishaps, an absurd absence of a schedule in the lives of the undisciplined students, and loosely proctored quizzes to restore the balance.",
    "The positive aspects of the situation include the ability to work around screen-sharing issues, the lack of a structured schedule for undisciplined students, and the opportunity to take advantage of lenient proctoring practices to restore balance.",
    "The benefits of the situation include the ability to adapt to screen-sharing challenges, the absence of a rigid schedule for undisciplined students, and the chance to take advantage of flexible proctoring policies to restore equilibrium.",
    "The advantages of the situation include the ability to work around screen-sharing difficulties, the lack of a structured schedule for undisciplined students, and the opportunity to take advantage of lenient proctoring practices to restore balance."
  ],
  [
    "There's no stringent bus driver who's gonna leave you behind if you're not at the stop on time; no way for the teacher to genuinely threaten you for your disrespectful actions.",
    "There is no strict bus driver or teacher who will penalize you for being late, and there are no consequences for disrespecting authority figures.",
    "There is no strict bus driver or teacher who will punish you for being late, and there are no repercussions for disregarding the rules of authority figures.",
    "There is no strict bus driver or teacher who will penalize you for being late, and there are no consequences for disobeying the rules of authority figures."
  ],
  [
    "Teachers have gone from the magicians in their five-star, one-person shows to street performers, begging for the least possible glance of interest at a medieval caf\u00e9.",
    "Teachers have gone from being masters of illusion and spectacle in their elaborate performances to struggling street performers, hoping for even a fleeting glance of interest from passersby at a modest medieval caf\u00e9.",
    "The once-glamorous world of magic and illusion has been replaced by the humblest of professions, as teachers now struggle to make ends meet and eke out a living in the most unassuming of settings.",
    "From the grand stages of their former careers to the cramped quarters of a medieval caf\u00e9, teachers have seen their fortunes plummet, leaving them with little more than the hope of a fleeting glance from a passerby as they struggle to make ends meet."
  ],
  [
    "There's no way they can suspend you, the only useful arrow in their otherwise hollow quavers, for being late to the zoom meeting; why of course, you had connectivity issues, who doesn't when there's a biology class at 8:30 in the morning?",
    "There is no way that teachers can be suspended for being late to a virtual meeting, as they have no other useful arrows in their quiver besides their ability to connect with students online.",
    "Despite the challenges of working remotely and dealing with connectivity issues, teachers are able to maintain their presence and continue to provide support to their students through virtual meetings.",
    "While it may be frustrating for teachers to experience technical difficulties during a meeting, they remain committed to maintaining their connection with students and providing them with the support they need to succeed."
  ],
  [
    "Another problem is the access to anonymity on specific platforms.",
    "Another issue that arises when using certain online platforms is the lack of anonymity for users, making it difficult for individuals to express themselves freely and without fear of repercussion.",
    "The use of certain online platforms can be problematic due to the lack of anonymity, as users may feel constrained in expressing their opinions or ideas without fear of backlash or retaliation.",
    "When using specific online platforms, users may experience difficulty in expressing themselves freely and authentically, due to the lack of anonymity provided by these platforms."
  ],
  [
    "Grant the class-clown a Guy-Fawkes curtain and suddenly, his jokes aren't gonna be quite PG13.",
    "Granting the class clown a platform for their jokes may result in them becoming more vulgar or inappropriate, as they are given greater freedom to express themselves without fear of censorship or repercussion.",
    "When the class clown is granted a platform for their jokes, they may become more irreverent and inappropriate, as they are able to express themselves without fear of censorship or repercussion.",
    "Allowing the class clown to have a platform for their jokes can result in them becoming more vulgar or inappropriate, as they are given greater freedom to express themselves without fear of censorship or repercussion."
  ]  
]
DELIMS = ('.','?','!')
class Doc:
    def __init__(self,path,id=None) -> None:
        self.path = path
        self.docdb = None #TODO connect to a sqlite storage of doc path, uuids.
        if(id is None):
            #TODOfetch uuid from sqlite or make entry in sqlite on new uuid.
            id = uuid.UUID(int=random.getrandbits(128)).hex
        self.id = id
class ParaData:
    def __init__(self,sentCount,style) -> None:
        self.sentCount = sentCount
        self.style = style
def getDataMembers(object):
    members = inspect.getmembers(object)
    members = filter(lambda member:not member[0].startswith('_'),members)
    members = filter(lambda member:not inspect.ismethod(member[1]),members)
    return list(members)
class DocVariationsGenerator:
    def __init__(self,alpha):
        self.picklepath = 'docvars.pkl'
        self.numphrases = 3
        self.nlp = None
        self.rephraser = None #initialized only if rephrasing function is called
        self.messageCharSet = []
        if(not alpha):
            for i in range(0,10):
                self.messageCharSet.append(str(i))
        for i in range(0,26):
            self.messageCharSet.append(chr(ord('a')+i))
    def extractSentences(self,doc:Doc):
        msDoc = Document(doc.path)
        paras = msDoc.paragraphs
        result = []
        paraSentCount:list[ParaData] = []
        for para in paras:
            paraSentCount.append(ParaData(0,para.style))
            para.text = para.text.strip()
            # print(para.text)
            sents = self.delimSplit(para.text)            
            paraSentCount[-1].sentCount = len(sents)
            result.extend(sents)
        return result,paraSentCount
    def injectSentences(self,doc:Doc,sentences,paraSentCount:list[ParaData]):
        msDoc = Document()
        pos = 0
        for paraData in paraSentCount:
            # print('sentcount: ',paraData.sentCount,'style: ',paraData.style.name)
            # print(getDataMembers(paraData.style.font))
            count = paraData.sentCount
            sents = sentences[pos:pos+count]
            paraText = ''
            if(len(sents)>0):
                paraText = ' '.join(sents)          
            msDoc.add_paragraph(paraText,paraData.style)
            pos += count
        msDoc.save(doc.path)
    def getCapacity(self,variations):
        capacity = 1
        for sentOptions in variations:
            capacity *= len(sentOptions)
        capacity = math.log2(capacity)
        return capacity#bits storeable in variations.
    def generateVariations(self, sentences,useDefaultVars=False):
        if(useDefaultVars):
            return DEFAULT_DOC_VARS
        if(self.rephraser is None):
            self.rephraser = Rephraser()
        variations = self.rephraser.paraphraseSentences(sentences,self.numphrases,minLength=5)
        return variations
    def pruneVariations(self,variations:list[list[str]]):
        for options in variations:
            if(len(options)>1):
                for i in range(1,len(options)):
                    delim = options[0][-1]
                    if(options[i][-1]!=delim):
                        if(not (options[i][-1] in DELIMS)):
                            options[i]+=delim
        #TODO discard sentences that are too far.
        if(self.nlp is None):
            self.nlp = spacy.load('en_core_web_lg')
        simils = []
        for j in range(len(variations)):
            options = variations[j]
            similarities = [1]
            filteredOptions = [options[0]]
            if(len(options)>1):
                truth = self.nlp(options[0])
                for i in range(1,len(options)):
                    option = self.nlp(options[i])
                    if(truth.similarity(option)>SIMILARITY_THRESHOLD):
                        filteredOptions.append(options[i])
            variations[j] = filteredOptions
        # print(simils)
        return variations
    def getFullDocVariations(self,doc:Doc):
        sentences,paraSentCount = self.extractSentences(doc)
        variations = self.generateVariations(sentences,True) # set to true to bypass GPT running locally.
        cleanVariations = self.pruneVariations(variations)
        return cleanVariations,paraSentCount
    def writePDF(self,sentences,name):
        # sents to docx.
        # try applying style.
        # to pdf.
        # use ghostscript
        pass
        # convert("input.docx")
        # convert("input.docx", "output.pdf")
        # convert("my_docx_folder/")
    def multibaseEncode(self,bases,message):
        ans = [0]*(len(bases)-1)
        for i in range(len(bases)-2,-1,-1):
            b = bases[i]
            ans[i] = math.floor(message/b)
            message = message - ans[i]*b
            if(message==0):
                break
        return ans
    def multibaseDecode(self,bases,indices):
        ans = 0
        for (index,base) in zip(indices,bases):
            ans += base*index
        return ans
    def getOffsets(self,cleanVars):
        offsets = [1]
        for sentenceOptions in cleanVars:
            offsets.append(len(sentenceOptions)*offsets[-1])
        return offsets
    def encodeIntoDoc(self,doc,message,outputFname='output.docx'):
        cleanVariations,paraSentCount = self.getFullDocVariations(doc)
        response = {
            "success":False
        }
        offsets = self.getOffsets(cleanVariations)
        capacity = offsets[-1] # last offset.
        if(message > capacity):
            response["error"] = f"Message ({message}) exceeds capacity ({capacity})."
            return response
        #message <= capacity
        indices = self.multibaseEncode(offsets,message)
        outputSentences=[]
        for index,sentenceOptions in zip(indices,cleanVariations):
            outputSentences.append(sentenceOptions[index])
        outputDoc = Doc(outputFname)
        self.injectSentences(outputDoc,outputSentences,paraSentCount)
        response["success"] = True
        return response
    def delimSplit(self,bulktext:str):
        regex = '|'.join('(?<={})'.format(re.escape(delim)) for delim in DELIMS)
        # print(regex)
        output = re.split(regex,bulktext)
        ans = []
        for i in range(len(output)):
            if(len(output[i])>0):
                ans.append(output[i].strip())
        return ans
    def decodeFromDoc(self,doc,outputDoc):
        cleanVariations,_ = self.getFullDocVariations(doc)
        sentences,_ = self.extractSentences(outputDoc)
        offsets = self.getOffsets(cleanVariations)
        indices = [0]*len(offsets)
        for i in range(0,len(sentences)):
            # try:
            indices[i] = cleanVariations[i].index(sentences[i])
            # except:
            # TODO: gracefully handle possible errors.  
            #     # print(cleanVariations[i])
            #     # print(sentences[i])
            #     # print(cleanVariations[0],cleanVariations[-1])
            #     # print(sentences[0],sentences[-1])
            #     # print(len(cleanVariations),len(sentences))
            #     return ''
        message = self.multibaseDecode(offsets,indices)
        return message
    def messageToNum(self,message:str):
        # allows for lowercase alphanum, spaced messages.
        ans = 0
        for i in range(0,len(message)):
            c = message[i]
            ans *= len(self.messageCharSet)
            ans += self.messageCharSet.index(c)
        return ans
    def numToMessage(self,num:int):
        message = ''
        l = len(self.messageCharSet)
        while(num>0):
            message += self.messageCharSet[num % l]
            num = math.floor(num/l)
        message = message[::-1]
        return message
    def customEncode(self,message:str)->bytearray:
        ans = []
        for i in range(0,len(message)):
            c = message[i]
            ans.append(self.messageCharSet.index(c))
        return bytearray(ans)
    def customDecode(self,message:bytearray)->str:
        ans = ''
        for b in message:
            ans += self.messageCharSet[b]
        return ans
    # def messageToNum2(self,message:str):
    #     num = 0
    #     compressedMessage = zlib.compress(self.customEncode(message))
    #     print(compressedMessage,len(compressedMessage))
    #     decompressed = zlib.decompress(compressedMessage) 
    #     print(self.customDecode(decompressed))
    #     return num
    # def numToMessage2(self,num:int):
    #     message = ''
    #     return message
def safeAttachExtension(filename:str,ext:str):
    if(not filename.endswith(ext)):
        # if(filename.endswith('.doc')):
        #     filename+='x'
        # else:
        filename+=ext
    return filename
if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--plain',type=str,help='Path to input .docx file.',required=True)
    parser.add_argument('--output',type=str,help='Path to output .docx file.',required=True)
    parser.add_argument('--message',type=str,help='Message to encode into input to produce output.')
    parser.add_argument('--decode',type=bool,help='Use this (--decode True) to decode output document using plain document',default=False)
    parser.add_argument('--alpha',type=bool,help='Use this to encode only alphabetic messages with higher capacity',default=False)
    args = parser.parse_args()
    dvg = DocVariationsGenerator(args.alpha)
    outFname:str = args.output
    doc:Doc = Doc(args.plain)
    
    if(args.decode):
        outputDoc:Doc = Doc(args.output)
        numMsg = dvg.decodeFromDoc(doc,outputDoc)
        message = dvg.numToMessage(numMsg)
        print(message)
    else:
        if(args.message is None):
            print('No message passed for encoding.')
            parser.print_help()
            exit()
        message = dvg.messageToNum(args.message)
        response = dvg.encodeIntoDoc(doc,message,outFname)
        if(response["success"]):
            print('Encoded doc saved to: ',outFname)
        else:
            print(response["error"])